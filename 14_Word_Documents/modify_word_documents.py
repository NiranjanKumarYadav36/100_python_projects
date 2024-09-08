import docx

doc_path1 = "pythonProjects/14_Word_Documents/panda1.docx"
doc_path2 = "pythonProjects/14_Word_Documents/panda2.docx"

doc1 = docx.Document(doc_path1)
doc2 = docx.Document(doc_path2)

# paragraphs = doc2.paragraphs[0].text
new_para = doc2.paragraphs[0]

paragraph = doc1.paragraphs
paragraph[1]._element.addnext(new_para._element)

doc1.save("14_Word_Documents/panda.docx")
